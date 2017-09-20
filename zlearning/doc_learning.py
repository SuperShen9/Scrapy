# -*- coding: utf-8 -*-

from docx import Document
doc=Document()

par = doc.add_paragraph('python is so easy')

doc.add_heading('about python')

doc.add_heading('驻足五秒',level=2)

table = doc.add_table(rows=2, cols=3)
cell = table.cell(0, 1)
cell.text = 'learn'
row=table.rows[1]
row.cells[0].text = 'first'
row.cells[1].text = 'second'


for row in table.rows:
    for cell in row.cells:
        print(cell.text)

row_count=len(table.rows)
col_count=len(table.columns)

print(row_count,col_count)

doc.add_paragraph('python',style='ListBullet')
par2 = doc.add_paragraph('learning')
par2.style='ListBullet'

par3=doc.add_paragraph()
par3.add_run('easy').bold=True

par4=doc.add_paragraph('Normal text | ')
par4.add_run('adding text','Emphasis')


doc.save('doc_text.docx')