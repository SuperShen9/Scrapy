# -*- coding: utf-8 -*-
import random,os
os.chdir('D:\zlianxi\shijuan')
import pprint
from docx import Document
for i in range(1,10):
    doc=Document()
    doc.add_heading('试卷'+str(i)+'题目:')
    dict={}
    list1 = ['+', '-']
    while True:
        a = random.randint(10, 99)
        b = random.randint(10, 99)
        c=random.choice(list1)
        if (c=='-' and a>b) or (c=='+' and a+b<100):
            astr=str(a)+c+str(b)
            bstr=str(b)+c+str(a)
            if dict.get(bstr,0)==0:
                dict.setdefault(astr,'aaa')
                if len(dict)==36:
                    break
    count1=0
    list2 = []
    for ky in dict.keys():
        list2.append(ky)
    table = doc.add_table(rows=9, cols=4)
    for k in range(9):
        for kk in range(4):
                cell = table.cell(k, kk)
                ky=list2[count1]
                cell.text=str(ky[:2])+ str(ky[2]) + str(ky[3:])+ ' ='+' '*10
                count1+=1

    doc.add_paragraph('\n')
    doc.add_heading('试卷'+str(i)+'答案:',level=2)
    doc.add_paragraph('-' *100)
    count2=0
    table1 = doc.add_table(rows=9, cols=4)
    for k in range(9):
        for kk in range(4):
            ky2 = list2[count2]

            cell = table1.cell(k, kk)
            if ky2[2] == '+':
                cell.text=str(int(ky2[:2]) + int(ky2[3:])).ljust(18)
            else:
                cell.text =str(int(ky2[:2]) - int(ky2[3:])).ljust(18)
            count2 += 1
    doc_name=str('试卷')+str(i)+'.docx'
    doc.save(doc_name)

